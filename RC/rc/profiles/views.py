from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.models import User, BaseUserManager
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.core.mail import send_mail
from .forms import ProfileForm
from docx import Document
import json, requests

# Create your views here.

host = 'http://192.168.0.28:8000/'

def agreement(request):
    f = None
    context = {}
    
    try:
        f = open('static/doc/개인정보 수집·이용 동의.docx', 'rb')
    except FileNotFoundError:
        context['agreement1'] = '개인정보 수집·이용 약관을 불러오는데 실패했습니다.'
    else:
        doc = Document(f)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        context['agreement1'] = '\n'.join(fullText)
    finally:
        if not (f is None):
            f.close()


    try:
        f = open('static/doc/개인정보 처리 동침 동의.docx', 'rb')
    except FileNotFoundError:
        context['agreement2'] = '개인정보 처리 동침 동의 약관을 불러오는데 실패했습니다.'
    else:
        doc = Document(f)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        context['agreement2'] = '\n'.join(fullText)
    finally:
        if not (f is None):
            f.close()

    return render(request, 'registration/agreement.html', context)


def getContext(user_id):
    user = get_object_or_404(User, pk=user_id)
    context = {
        "username" : user.username,
        "email" : user.email,
        "image" : user.profile.image,
        "gender" : user.profile.gender,
        'birth_year': user.profile.birth_year,
        'birth_month': user.profile.birth_month,
        'birth_date': user.profile.birth_date,
        "type" : user.profile.type,
        "status" : user.profile.status
    }
    return context


def done(request):
    return render(request, 'registration/done.html')


def account_edit(request, account_id=None):
    context = {}

    if request.method == 'POST':
        form = ProfileForm(request.POST, request.FILES)

        if form.is_valid():
            user = form.save()
            context['title'] = '회원가입 완료'
            context['messages'] = ['환영합니다.', '메인화면으로 이동합니다.', '로그인을 해주세요.']
        else:
            user = get_object_or_404(User, pk=account_id)
            user.set_password(form.cleaned_data.get('password1'))
            user.email = form.cleaned_data.get('email')
            # update_session_auth_hash(request, user)
            context['title'] = '수정되었습니다.'
            context['messages'] = ['메인화면으로 이동합니다.', '다시 로그인을 해주세요.']
        user.profile.gender = form.cleaned_data.get('gender')
        user.profile.birth_year = form.cleaned_data.get('birth_year')
        user.profile.birth_month = form.cleaned_data.get('birth_month')
        user.profile.birth_date = form.cleaned_data.get('birth_date')
        user.save()
        
        target = get_object_or_404(User, username=request.POST.get("username"))
        url = host + "init_wallet/" + str(target.pk)
        res = requests.get(url)
        if res == "fail":
            context['messages'] = ['계좌 생성에 실패했습니다.', '관리자에게 문의하세요.', '메인화면으로 이동합니다.', '로그인을 해주세요.']
        
        return render(request, 'registration/done.html', context)
    else:
        if account_id:
            context = getContext(account_id)
            return render(request, 'registration/profile_edit.html', context)
        else:
            form = ProfileForm()
            return render(request, 'registration/profile_edit.html', { 'form' : form })


def my_info(request, account_id=None):
    template_name = "registration/profile_myInfo.html"
    context = getContext(account_id)
    url = host +"get_account/" + str(account_id)
    response = requests.get(url)
    res = json.loads(response.text)
    data = {
        "context" : context,
        "balance" : res['value']
    }
    return render(request, template_name, data)



def getUserList(request):
    user_queryset = User.objects.filter(email=request.GET.get("email"))
    user_list = []
    for user in user_queryset:
        user_list.append(user.username)
    data = {
        'user_list' : user_list
    }
    json_data = json.dumps(data)
    return HttpResponse(json_data, content_type="application/json;charset=UTF-8")


def initPassword(request):
    username = request.GET.get("username")
    email = request.GET.get("email")
    user = User.objects.filter(username=username, email=email)
    if user:
        # char_set = 'abcdefghjkmnpqrstuvwxyzABCDEFGHJKLMNPQRSTUVWXYZ23456789'
        # new_pwd = BaseUserManager().make_random_password(length=10, allowed_chars=char_set)
        # user[0].set_password(new_pwd)
        # user[0].save()
        # send_mail(
        #     '[RC: 발신전용] 비밀번호 안내 메일',
        #     '비밀번호가 [ ' + new_pwd + ' ]로 초기화 되었습니다.',
        #     'doradora46@naver.com',
        #     [user[0].email],
        #     fail_silently=False,
        # )
        data = { 'result' : True }
    else:
        data = { 'result' : False }
    json_data = json.dumps(data)
    return HttpResponse(json_data, content_type="application/json;charset=UTF-8")


def search_username(request):
    return render(request, 'registration/search_username.html')


def search_password(request):
    return render(request, 'registration/search_password.html')


def check_username(request):
    username = request.GET.get('username', None)
    res = User.objects.filter(username=username).exists()
    data = {
        'result' : res
    }
    json_data = json.dumps(data)
    return HttpResponse(json_data, content_type="application/json;charset=UTF-8")


def check_password(request, account_id=None):
    context = {}
    if request.method == 'POST':
        user = request.user.check_password(request.POST.get('password'))
        print(request.POST.get('password'))
        print(user)
        if user:
            return redirect('profile:account_edit', account_id=account_id)
        else:
            context['error'] = 'true'
    return render(request, 'registration/check_password.html', context)


# --- @login_required
class LoginRequiredMixin(object):
    @classmethod
    def as_view(cls, **initkwargs):
        view = super(LoginRequiredMixin, cls).as_view(**initkwargs)
        return login_required(view)

